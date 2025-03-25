from flask import Flask, request, jsonify
from sentence_transformers import SentenceTransformer
from transformers import BartForConditionalGeneration, BartTokenizer
import pdfplumber
import docx
import re
import spacy
import numpy as np
import jwt
import datetime
import logging
import os
from pymongo import MongoClient
from flask_cors import CORS
from functools import wraps
from bson import ObjectId

# Initialize Flask app
app = Flask(__name__)
CORS(app, supports_credentials=True, origins=["http://localhost:3000"])

# Load environment variables
SECRET_KEY = os.getenv("JWT_SECRET_KEY", "default_secret_key")

# Set up logging
logging.basicConfig(level=logging.INFO)

# Connect to MongoDB
try:
    mongo_client = MongoClient("mongodb://localhost:27017/")
    db = mongo_client["DocAssist"]
    documents_collection = db["documents"]
    logging.info("Connected to MongoDB successfully!")
except Exception as e:
    logging.error(f"Error connecting to MongoDB: {e}")
    exit(1)

# Load NLP models
try:
    bart_model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn").to("cpu")
    bert_model = SentenceTransformer('all-MiniLM-L6-v2', device="cpu")
    bart_tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    logging.error(f"Error loading models: {e}")
    exit(1)

# Authentication middleware
def auth_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get("Authorization")

        if token:
            token_parts = token.split(" ")
            if len(token_parts) == 2 and token_parts[0] == "Bearer":
                token = token_parts[1]
            else:
                return jsonify({"error": "Invalid Authorization format!"}), 401
        else:
            token = request.cookies.get("token")
            if not token:
                return jsonify({"error": "Authentication token is missing!"}), 401

        try:
            decoded_token = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
            request.user_id = decoded_token.get("user_id")
            if not request.user_id:
                return jsonify({"error": "Invalid token: user_id missing!"}), 401
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token has expired!"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Invalid token!"}), 401

        return f(*args, **kwargs)
    return decorated_function

# Helper functions
def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

def clean_text(text):
    return re.sub(r"\s+", " ", text).strip()

def split_into_chunks(text, max_chunk_size=512):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    chunks, current_chunk = [], ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chunk_size:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

def cosine_similarity(vec1, vec2):
    return np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))

@app.route("/")
def home():
    return "Welcome to the DocAssist API!"

# API: Process document
@app.route("/process-document", methods=["POST"])
@auth_required
def process_document():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file provided!"}), 400

        logging.info(f"Processing file: {file.filename}")

        if file.filename.endswith(".txt"):
            content = file.read().decode("utf-8")
        elif file.filename.endswith(".pdf"):
            content = extract_text_from_pdf(file)
        elif file.filename.endswith(".docx"):
            content = extract_text_from_docx(file)
        else:
            return jsonify({"error": "Unsupported file format!"}), 400

        cleaned_content = clean_text(content)
        if not cleaned_content:
            return jsonify({"error": "Extracted content is empty!"}), 400

        chunks = split_into_chunks(cleaned_content)
        embeddings = [{"text": chunk, "embedding": bert_model.encode(chunk).tolist()} for chunk in chunks]

        if not embeddings:
            return jsonify({"error": "No valid embeddings generated!"}), 400

        document_id = str(ObjectId())

        document_entry = {
            "user_id": request.user_id,
            "document_id": document_id,
            "document_name": file.filename,
            "document": cleaned_content,
            "embeddings": embeddings,
            "timestamp": datetime.datetime.utcnow()
        }

        documents_collection.insert_one(document_entry)
        logging.info("Document successfully saved in MongoDB")

        return jsonify({"message": "File processed and embedded successfully!", "document_id": document_id})

    except Exception as e:
        logging.error(f"Error processing document: {str(e)}")
        return jsonify({"error": f"Error processing document: {str(e)}"}), 500

# API: Get suggestions
@app.route("/get-suggestions", methods=["POST"])
@auth_required
def get_suggestions():
    try:
        data = request.json
        query = data.get("query", "")

        if not query:
            return jsonify({"error": "Query is required!"}), 400

        user_documents = list(documents_collection.find({"user_id": request.user_id}))
        if not user_documents:
            return jsonify({"error": "No documents found for the user!"}), 404

        query_embedding = bert_model.encode(query)
        ranked_chunks = []

        for doc in user_documents:
            for chunk in doc.get("embeddings", []):
                similarity = cosine_similarity(query_embedding, chunk["embedding"])
                ranked_chunks.append((chunk["text"], similarity))

        ranked_chunks.sort(key=lambda x: x[1], reverse=True)
        top_chunks = " ".join(chunk[0] for chunk in ranked_chunks[:5])

        inputs = bart_tokenizer.encode(top_chunks, return_tensors="pt", max_length=1024, truncation=True)
        summary_ids = bart_model.generate(inputs, max_length=300, min_length=100, length_penalty=2.0, num_beams=4, early_stopping=True)
        summary = bart_tokenizer.decode(summary_ids[0], skip_special_tokens=True)

        return jsonify({"summary": summary}), 200

    except Exception as e:
        logging.error(f"Error generating suggestions: {str(e)}")
        return jsonify({"error": "Error generating suggestions"}), 500

if __name__ == "__main__":
    app.run(debug=False, port=5002)









from flask import Flask, request, jsonify
from sentence_transformers import SentenceTransformer
from transformers import BartForConditionalGeneration, BartTokenizer
import pdfplumber
import docx
import re
import spacy
import numpy as np
import logging
from flask_cors import CORS

# Initialize the Flask app and CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
# Set up logging
logging.basicConfig(level=logging.INFO)

# Root route
@app.route("/")
def home():
    return jsonify({"message": "Welcome to the Flask server!"})

# Load models
bart_model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn").to("cpu")
bert_model = SentenceTransformer('all-MiniLM-L6-v2', device="cpu")
bart_tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
nlp = spacy.load("en_core_web_sm")

# Helper functions
def clean_text(text):
    return re.sub(r"\s+", " ", text).strip()

def split_into_chunks(text, max_chunk_size=512):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    chunks, current_chunk = [], ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chunk_size:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def process_text(content):
    cleaned_content = clean_text(content)
    chunks = split_into_chunks(cleaned_content)
    embeddings = bert_model.encode(chunks)
    return [{"text": chunk, "embedding": embeddings[i].tolist()} for i, chunk in enumerate(chunks)]

def process_pdf(file):
    with pdfplumber.open(file) as pdf:
        content = "".join(page.extract_text() for page in pdf.pages)
    return process_text(content)

def process_docx(file):
    doc = docx.Document(file)
    content = "\n".join(para.text for para in doc.paragraphs)
    return process_text(content)

def get_suggestions(user_input, document_chunks):
    query_embedding = bert_model.encode(user_input)
    
    def cosine_similarity(a, b):
        return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
    
    results = sorted(
        [(chunk["text"], cosine_similarity(query_embedding, chunk["embedding"])) for chunk in document_chunks],
        key=lambda x: x[1],
        reverse=True
    )[:3]
    
    suggestions = []
    for i, (text, score) in enumerate(results):
        input_ids = bart_tokenizer.encode(text, return_tensors='pt', max_length=1024, truncation=True)
        summary_ids = bart_model.generate(input_ids, max_length=300, min_length=100, length_penalty=2.0, num_beams=4, early_stopping=True)
        summarized_text = bart_tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        
        # Remove bullet points and format as plain text
        plain_text = summarized_text.replace("\n", " ").replace("•", "").strip()
        suggestions.append({"id": i + 1, "text": plain_text})
    
    return suggestions

# API Endpoints
@app.route("/process-document", methods=["POST"])
def process_document():
    try:
        document_chunks = []
        file = request.files.get("file")
        
        if not file:
            return jsonify({"error": "No file provided!"}), 400
        
        logging.info(f"Processing file: {file.filename}")
        
        if file.filename.endswith(".txt"):
            content = file.read().decode("utf-8")
            document_chunks = process_text(content)
        elif file.filename.endswith(".pdf"):
            document_chunks = process_pdf(file)
        elif file.filename.endswith(".docx"):
            document_chunks = process_docx(file)
        else:
            return jsonify({"error": "Unsupported file format!"}), 400
        
        logging.info(f"Processed {len(document_chunks)} chunks.")
        return jsonify({"message": "File processed successfully!", "chunks": document_chunks})

    except Exception as e:
        logging.error(f"Error processing document: {str(e)}")
        return jsonify({"error": f"Error processing document: {str(e)}"}), 500

@app.route("/get-suggestions", methods=["POST"])
def get_suggestions_api():
    try:
        data = request.json
        user_input = data.get("user_input", "")
        document_chunks = data.get("document_chunks", [])

        if not user_input or not document_chunks:
            return jsonify({"error": "Missing user input or document chunks!"}), 400

        logging.info(f"User input: {user_input}")
        logging.info(f"Number of document chunks: {len(document_chunks)}")

        suggestions = get_suggestions(user_input, document_chunks)
        logging.info(f"Suggestions generated: {suggestions}")

        return jsonify(suggestions)

    except Exception as e:
        logging.error(f"Error generating suggestions: {str(e)}")
        return jsonify({"error": f"Error generating suggestions: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True,port=5000)